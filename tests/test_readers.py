import pytest
from whisper_notes.readers import read_document


def test_read_txt(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("Hello world", encoding="utf-8")
    assert read_document(f) == "Hello world"


def test_read_md(tmp_path):
    f = tmp_path / "test.md"
    f.write_text("# Title\n\nContent", encoding="utf-8")
    assert read_document(f) == "# Title\n\nContent"


def test_read_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    f = tmp_path / "test.docx"
    doc.save(str(f))
    result = read_document(f)
    assert "First paragraph" in result
    assert "Second paragraph" in result


def test_read_rtf(tmp_path):
    f = tmp_path / "test.rtf"
    f.write_text(r"{\rtf1 Hello RTF world}", encoding="utf-8")
    result = read_document(f)
    assert "Hello RTF world" in result


def test_read_unsupported(tmp_path):
    f = tmp_path / "test.pages"
    f.touch()
    with pytest.raises(ValueError, match="Unsupported"):
        read_document(f)
